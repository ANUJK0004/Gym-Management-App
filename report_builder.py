from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'SweatSync_Project_Report_Updated.docx'
ASSETS = ROOT / 'report_assets'
ASSETS.mkdir(exist_ok=True)

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
LIGHT = 'F2F4F7'
MID = 'E8EEF5'


def make_dfd(path, level):
    image = Image.new('RGB', (1800, 1120), 'white'); draw = ImageDraw.Draw(image)
    font = ImageFont.truetype('C:/Windows/Fonts/arial.ttf', 28); small = ImageFont.truetype('C:/Windows/Fonts/arial.ttf', 21); title_font = ImageFont.truetype('C:/Windows/Fonts/arialbd.ttf', 36)
    def xy(x, y): return int(x*180), int((6.2-y)*180)
    def centered_text(rect, text, f=font, fill='#0B2545'):
        lines=text.split('\n'); total=sum(draw.textbbox((0,0),line,font=f)[3] for line in lines)+8*(len(lines)-1); cy=(rect[1]+rect[3]-total)//2
        for line in lines:
            b=draw.textbbox((0,0),line,font=f); draw.text(((rect[0]+rect[2]-(b[2]-b[0]))//2,cy),line,font=f,fill=fill); cy+=b[3]-b[1]+8
    def box(x, y, w, h, text, color='#E8EEF5', radius=.12):
        left,bottom=xy(x,y); right,top=xy(x+w,y+h); rect=(left,top,right,bottom); draw.rounded_rectangle(rect,radius=int(radius*180),fill=color,outline='#1F4D78',width=3); centered_text(rect,text)
    def store(x,y,w,h,text):
        left,bottom=xy(x,y); right,top=xy(x+w,y+h); rect=(left,top,right,bottom); draw.rectangle(rect,fill='white',outline='#1F4D78',width=3); draw.line((left+20,top+20,right-20,top+20),fill='#1F4D78',width=2); centered_text(rect,text,small)
    def arrow(a,b,label='',offset=0):
        x1,y1=xy(*a); x2,y2=xy(*b); draw.line((x1,y1,x2,y2),fill='#455A64',width=3); import math; ang=math.atan2(y2-y1,x2-x1); size=15; draw.polygon([(x2,y2),(x2-size*math.cos(ang-.5),y2-size*math.sin(ang-.5)),(x2-size*math.cos(ang+.5),y2-size*math.sin(ang+.5))],fill='#455A64')
        if label:
            bnd=draw.textbbox((0,0),label,font=small); mx=(x1+x2-(bnd[2]-bnd[0]))//2; my=(y1+y2-(bnd[3]-bnd[1]))//2-int(offset*180); draw.rectangle((mx-4,my-3,mx+(bnd[2]-bnd[0])+4,my+(bnd[3]-bnd[1])+3),fill='white'); draw.text((mx,my),label,font=small,fill='#37474F')
    if level == 0:
        box(.35,4.5,1.45,.72,'OWNER', '#F4F6F9'); box(.35,1.0,1.45,.72,'MEMBER', '#F4F6F9')
        box(4.0,2.35,2.1,1.15,'SWEATSYNC\nFLUTTER APPLICATION','#DCEAF7')
        box(8.05,4.5,1.55,.72,'FIREBASE\nAUTH', '#F4F6F9'); store(7.9,2.65,1.8,.72,'CLOUD\nFIRESTORE'); box(8.05,.9,1.55,.72,'CLOUD FUNCTIONS\n& STORAGE','#F4F6F9')
        arrow((1.8,4.86),(4,3.1),'management requests',.18); arrow((4,3.45),(1.8,5.1),'dashboards / results',.18)
        arrow((1.8,1.35),(4,2.7),'profile, membership, workout data',-.22); arrow((4,2.45),(1.8,1.1),'member experience',-.2)
        arrow((6.1,3.1),(8.05,4.86),'sign-in / role',.15); arrow((8.05,4.55),(6.1,3.35),'identity state',.14)
        arrow((6.1,2.9),(7.9,3.0),'read / write',.16); arrow((7.9,2.68),(6.1,2.6),'documents / analytics',-.15)
        arrow((6.1,2.55),(8.05,1.25),'privileged enrollment / export',-.2); arrow((8.05,1.6),(6.1,2.35),'validated result / download URL',.22)
        title='Figure 1. Context-level DFD for SweatSync'
    else:
        box(.25,4.95,1.28,.65,'OWNER', '#F4F6F9'); box(.25,.65,1.28,.65,'MEMBER', '#F4F6F9')
        box(2.05,4.55,1.75,.85,'1.0 Authentication\n& Role Access','#DCEAF7')
        box(4.25,4.55,1.75,.85,'2.0 Owner\nOperations','#DCEAF7')
        box(4.25,.55,1.75,.85,'3.0 Member\nExperience','#DCEAF7')
        box(6.5,2.55,1.85,.85,'4.0 Finance\n& Reports','#DCEAF7')
        store(8.55,4.35,1.2,.7,'USERS /\nGYMS'); store(8.55,2.55,1.2,.7,'PLANS /\nENROLLMENTS'); store(8.55,.75,1.2,.7,'TRANSACTIONS /\nREPORT FILES')
        arrow((1.53,5.25),(2.05,5.0),'credentials',.14); arrow((3.8,5.0),(4.25,5.0),'authorized owner',.14); arrow((6.0,5.0),(8.55,4.7),'create / update',.14)
        arrow((1.53,1.0),(4.25,1.0),'member requests',.14); arrow((6.0,1.0),(8.55,1.1),'read / write',.14)
        arrow((6.0,4.8),(6.5,3.25),'finance / report request',.12); arrow((8.35,2.95),(8.55,2.95),'plan data',.12); arrow((8.35,2.75),(8.55,1.1),'export data',-.13)
        arrow((6.5,2.85),(1.53,5.0),'analytics / export link',-.15); arrow((4.25,1.25),(1.53,.9),'membership / workout status',-.15)
        title='Figure 2. Level-1 DFD for implemented SweatSync modules'
    b=draw.textbbox((0,0),title,font=title_font); draw.text(((1800-(b[2]-b[0]))//2,20),title,font=title_font,fill='#0B2545')
    image.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = mar.find(qn(f'w:{side}'))
        if node is None: node = OxmlElement(f'w:{side}'); mar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'),'dxa')

def set_table_widths(table, widths):
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); tblPr.append(layout)
    for row in table.rows:
        for cell, width in zip(row.cells, widths): cell.width = Inches(width)

def field(paragraph, instruction):
    r = paragraph.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin'); instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=instruction; separate=OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'),'separate'); text=OxmlElement('w:t'); text.text='1'; end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end'); r._r.extend([begin,instr,separate,text,end])

def add_page_break(doc): doc.add_page_break()

def add_para(doc, text='', bold_lead=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.1
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold=True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

def heading(doc, text, level=1): return doc.add_heading(text, level=level)

def bullets(doc, values):
    for value in values:
        p=doc.add_paragraph(style='List Bullet'); p.add_run(value)

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for c,h in zip(t.rows[0].cells,headers):
        c.text=h; set_cell_shading(c,LIGHT); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True
    trPr = t.rows[0]._tr.get_or_add_trPr(); tbl_header = OxmlElement('w:tblHeader'); tbl_header.set(qn('w:val'), 'true'); trPr.append(tbl_header)
    for row in rows:
        cells=t.add_row().cells
        for c,val in zip(cells,row):
            c.text=val; set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths: set_table_widths(t,widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def placeholder(doc, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run('\n\n\n\n\n\n\n\nINSERT GOOGLE PIXEL 10 EMULATOR SCREENSHOT HERE\n\n\n\n\n\n')
    run.font.color.rgb=RGBColor(110,110,110); run.font.size=Pt(10); run.bold=True
    p.paragraph_format.space_after=Pt(4)
    pPr=p._p.get_or_add_pPr(); borders=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'8'); bottom.set(qn('w:space'),'8'); bottom.set(qn('w:color'),'A6A6A6'); borders.append(bottom); pPr.append(borders)
    cap=doc.add_paragraph(caption); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs: r.italic=True; r.font.size=Pt(9)

def configure(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for style_name,size,color,before,after in [('Title',24,INK,0,8),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK_BLUE,8,4)]:
        st=doc.styles[style_name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('SweatSync Project Report | Page '); field(footer,'PAGE')

def build():
    make_dfd(ASSETS/'dfd_context.png',0); make_dfd(ASSETS/'dfd_level1.png',1)
    doc=Document(); configure(doc)
    # Cover
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SWEATSYNC'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(INK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SMART GYM MANAGEMENT AND MEMBER ENGAGEMENT PLATFORM'); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('USING FLUTTER, DART AND FIREBASE').italic=True
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('A Project Report').bold=True; p.runs[0].font.size=Pt(16)
    p=doc.add_paragraph('in partial fulfillment of the requirements for the award of the degree of'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph('B.Tech in Computer Science and Engineering'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
    p=doc.add_paragraph('with specialization in Artificial Intelligence and Machine Learning'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph('under'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph('Academy of Skill Development'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph('Submitted by'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph('Anuj Kushwah'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(14)
    p=doc.add_paragraph('Anand Engineering College, Keetham-Agra'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_page_break(doc)
    heading(doc,'CERTIFICATE',1)
    add_para(doc,'This is to certify that Anuj Kushwah has successfully carried out the project titled "SweatSync - Smart Gym Management and Member Engagement Platform" under my supervision. The work is submitted in partial fulfillment of the requirements for the award of the degree of B.Tech in Computer Science and Engineering with specialization in Artificial Intelligence and Machine Learning.')
    add_para(doc,'The report describes the current design, development, integration and testing of the Flutter application. It covers only the functionality currently present in the codebase and identifies unfinished work separately in the feature-improvement section.')
    doc.add_paragraph('\n\nDate: __________________\n\n______________________________\nSignature of the Mentor\n[ANINDYA MUKHERJEE]')
    add_page_break(doc); heading(doc,'ACKNOWLEDGEMENT',1)
    add_para(doc,'I take this opportunity to express my deep gratitude and sincerest thanks to my project mentor, ANINDYA MUKHERJEE, for valuable suggestions, helpful guidance and encouragement during this project work.')
    add_para(doc,'I am thankful to everyone who contributed feedback, testing, technical discussion and review. I also acknowledge the Flutter, Dart and Firebase ecosystems and the open-source libraries used in the application. Finally, I am grateful to the faculty members of the Academy of Skill Development for their support.')
    add_page_break(doc); heading(doc,'ABSTRACT OF THE PROJECT',1)
    add_para(doc,'SweatSync is a role-aware mobile application for gym operations and member engagement. It is built with Flutter and Dart, with Firebase Authentication for identity, Cloud Firestore for application data, Cloud Functions for privileged workflows and Cloud Storage for generated export files. The present implementation provides owner and member experiences; it also includes owner-managed trainer enrollment and management, while the dedicated trainer application is not yet implemented.')
    add_para(doc,'The owner experience includes gym setup, dashboard summaries, member management, membership-plan management, member enrollment, trainer management, finance, settings and reporting. The member experience includes role-aware access, profile setup and editing, membership information, assigned workouts, workout-session completion and progress views. Riverpod and GoRouter coordinate state and navigation in a feature-first architecture. Attendance and a complete trainer dashboard are not implemented and are therefore documented as future improvements rather than current features.')
    add_page_break(doc); heading(doc,'CONTENTS',1)
    for x in ['1. Introduction','2. Feasibility Study','3. Scope of the Project','4. Software Requirement Specification','5. Hardware and Software Requirements','6. Data Flow Diagrams','7. Feature-First Architecture, Coding Standards and Screenshots','8. Application Testing','9. Feature Improvement and Future Enhancements','10. Conclusion','11. References and Bibliography','Appendix A. Current Implementation Status']:
        add_para(doc,x)
    heading(doc,'1. INTRODUCTION',1)
    add_para(doc,'Gym operations often rely on paper registers, spreadsheets and separate communication channels. This makes it difficult to keep member records, membership plans, trainer records, payments and member activity consistent. SweatSync centralizes these workflows in a mobile application while presenting different navigation and data to owners and members.')
    add_para(doc,'The application uses a dark, mobile-first interface and separates application work into feature modules. Firebase-backed data and callable server functions provide controlled operations for sensitive enrollment and report-export workflows. The resulting structure is intended to remain maintainable as completed modules grow.')
    heading(doc,'2. FEASIBILITY STUDY',1)
    heading(doc,'2.1 Technical Feasibility',2); add_para(doc,'Flutter and Dart provide a single codebase for Android-focused mobile delivery. Firebase Authentication, Cloud Firestore, Cloud Functions and Cloud Storage supply managed services for the current client and server workflows. Riverpod, GoRouter and the feature-first structure keep state, navigation and persistence responsibilities separated.')
    heading(doc,'2.2 Operational Feasibility',2); add_para(doc,'Owners can access management, finance and report workflows from the owner shell, while members receive dashboard, membership, workout, progress and profile experiences. Forms, status chips, refresh behavior and controlled enrollment steps support routine gym operations on a phone.')
    heading(doc,'2.3 Economic Feasibility',2); add_para(doc,'The project uses Flutter and Firebase services, reducing the need to create and operate a separate custom application server for the current scope. Managed Firebase services are suitable for incremental development and controlled deployment of callable operations.')
    heading(doc,'3. SCOPE OF THE PROJECT',1)
    add_para(doc,'The present scope covers implemented owner and member workflows only. The application supports authentication and profile setup, gym information, membership-plan administration, member enrollment and management, trainer enrollment and management, finance, reporting, member membership details, workouts, workout completion, progress and profile handling.')
    table(doc,['Role','Implemented capabilities'],[
        ['Owner','Gym setup and editing; dashboard summaries; member search, filters and details; membership plans; member enrollment; trainer management; finance; reports; owner settings.'],
        ['Member','Role-aware entry; profile setup/editing; dashboard summary; membership details; today\'s workout; workout session and completion; progress view; profile.'],
        ['Trainer','Trainer enrollment and owner-side trainer management are present. The dedicated trainer shell currently displays a placeholder and is not a complete role experience.'],
    ],[1.1,5.4])
    heading(doc,'4. SOFTWARE REQUIREMENT SPECIFICATION',1)
    heading(doc,'4.1 Functional Requirements',2)
    bullets(doc,[
        'Authenticate users and route completed owner and member profiles to role-aware application experiences.',
        'Allow an owner to create and update gym information, manage members, and search/filter member records.',
        'Allow an owner to create, edit, activate, deactivate and delete membership plans with validation.',
        'Support a multi-step member enrollment flow using active plans, fitness goals, start dates and UPI, cash, card or pending payment selection.',
        'Allow an owner to manage trainer records, search trainers and use the controlled trainer-enrollment workflow.',
        'Present owner finance data, transactions, analytics and report-export workflows.',
        'Allow members to view membership status, access assigned workouts, complete workouts and view progress/profile data.'
    ])
    heading(doc,'4.2 Non-Functional Requirements',2)
    bullets(doc,['Usability: mobile-oriented screens, loading states, error states and refresh controls.','Security: Firebase Authentication and owner/gym validation for privileged callable operations.','Maintainability: presentation, domain and data layers remain separated inside features.','Reliability: membership-plan state and cross-gym relationships are validated before sensitive mutations.','Performance: asynchronous providers fetch data and refresh invalidated state after changes.'])
    heading(doc,'4.3 Technology Requirements',2)
    table(doc,['Component','Technology / purpose'],[
        ['Client','Flutter and Dart mobile application'],['State and navigation','flutter_riverpod and go_router'],['Identity and data','Firebase Core, Firebase Authentication and Cloud Firestore'],['Server workflows','Cloud Functions for member enrollment, trainer enrollment, report export and finance export'],['File storage','Firebase Storage / Cloud Storage for generated files'],['UI and data visualization','Flutter Material, custom design system and fl_chart'],['Supporting packages','mobile_scanner, connectivity_plus, flutter_secure_storage, Firebase Messaging and local notifications packages are declared in the project.']
    ],[1.55,4.95])
    heading(doc,'5. HARDWARE AND SOFTWARE REQUIREMENTS',1)
    table(doc,['Category','Requirement'],[
        ['Development machine','Windows 10 or later, 8 GB RAM minimum, stable internet connection and sufficient storage for Flutter/Android tooling.'],
        ['Testing device','Android phone or Android Emulator; the supplied environment includes Pixel 10 AVD profiles.'],
        ['Development tools','Flutter SDK, Dart SDK, Android Studio or VS Code, Android SDK/AVD Manager, JDK and Git.'],
        ['Backend configuration','Firebase project configuration for Authentication, Firestore, Cloud Functions and Storage.']
    ],[1.45,5.05])
    heading(doc,'6. DATA FLOW DIAGRAMS',1)
    add_para(doc,'A data flow diagram (DFD) shows how information moves between people, application processes and data stores. The diagrams below document the current Firebase-backed SweatSync scope; they do not represent planned attendance or a future trainer dashboard.')
    pic=doc.add_picture(str(ASSETS/'dfd_context.png'),width=Inches(6.35)); pic._inline.docPr.set('descr','Context-level data flow diagram for the SweatSync application.'); p=doc.add_paragraph('Figure 1. Context-level DFD for SweatSync.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pic=doc.add_picture(str(ASSETS/'dfd_level1.png'),width=Inches(6.35)); pic._inline.docPr.set('descr','Level-1 data flow diagram for implemented SweatSync modules.'); p=doc.add_paragraph('Figure 2. Level-1 DFD for implemented SweatSync modules.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc,'Owner and member actions reach Flutter screens, Riverpod providers and repositories. Standard data reads and writes use Firestore. Privileged member enrollment, trainer enrollment and export actions call Cloud Functions, which validate the authenticated caller and return a controlled result. Generated export files can be stored in Cloud Storage and opened through a download URL.')
    heading(doc,'7. FEATURE-FIRST ARCHITECTURE, CODING STANDARDS AND SCREENSHOTS',1)
    heading(doc,'7.1 Feature-First Architecture',2)
    add_para(doc,'As in the reference report, SweatSync follows a feature-first architecture. Instead of grouping the whole application only by technical type, each business capability is represented as a feature folder. Examples include auth, dashboard/member, dashboard/owner, member_management, membership_plan, trainer_management, finance, report, workout, progress, profile and gym. This makes related screens, providers, repositories, entities and data sources easier to locate and extend.')
    table(doc,['Layer','Responsibility','Examples'],[
        ['Presentation','Flutter screens, widgets and Riverpod providers that display and coordinate UI state.','MemberHomeScreen, OwnerHomeScreen, ReportScreen, provider classes.'],
        ['Domain','Business entities and repository contracts independent of Firebase UI details.','MembershipPlan, ManagedMember, Workout, FinanceTransaction.'],
        ['Data','Firestore/Cloud Function data sources, models and repository implementations.','Remote data sources, model classes, repository implementations.'],
        ['Backend','Privileged callable server operations and guarded mutations.','createMemberEnrollment, createTrainerEnrollment, exportReport, exportFinanceReport.']
    ],[1.0,2.2,3.3])
    heading(doc,'7.2 Coding Standards Followed',2)
    bullets(doc,['Feature-oriented folder naming and separation of presentation, domain and data responsibilities.','Repository interfaces in the domain layer and Firebase-specific implementations in the data layer.','Riverpod providers for asynchronous state, dependency wiring and explicit refresh/invalidation after updates.','GoRouter routes for application navigation and role-aware redirect decisions.','Reusable design-system components for cards, form fields, navigation, spacing, colors and buttons.','Explicit loading, empty and error UI states in key list, dashboard and report screens.'])
    heading(doc,'7.3 Screenshot Placeholders',2)
    add_para(doc,'A Pixel 10 emulator was available but could not reach a stable, authenticated application screen during this report update. The following labeled spaces are intentionally left for real screenshots to be inserted after a successful Firebase-backed run.')
    placeholder(doc,'Figure 3. Screenshot placeholder - Role selection / authentication screen.')
    placeholder(doc,'Figure 4. Screenshot placeholder - Owner dashboard or member-management screen.')
    placeholder(doc,'Figure 5. Screenshot placeholder - Member dashboard, workout or progress screen.')
    placeholder(doc,'Figure 6. Screenshot placeholder - Membership plan, trainer management, finance or reports screen.')
    heading(doc,'8. APPLICATION TESTING',1)
    heading(doc,'8.1 Testing Objectives',2); add_para(doc,'Testing verifies that actions produce the expected application state, that UI feedback is understandable and that restricted owner workflows are protected. The focus is on functional correctness, integration between Flutter and Firebase services, and user-facing loading, error and refresh behavior.')
    heading(doc,'8.2 Testing Strategy',2)
    bullets(doc,['Unit-level verification of entities, provider behavior and validation logic.','Integration checks for Flutter -> repository -> Firestore/Cloud Function workflows.','UI checks for forms, selection steps, success/error feedback and list refresh.','Regression checks after changes to related modules such as finance and exports.'])
    heading(doc,'8.3 Representative Test Cases',2)
    table(doc,['ID','Test condition','Expected result'],[
        ['TC-01','Owner opens member management.','Members for the owner\'s gym load with search/filter support.'],
        ['TC-02','Owner creates a unique membership plan.','The plan is saved and the list refreshes.'],
        ['TC-03','Owner disables a plan.','The plan becomes inactive and cannot be used for a new enrollment.'],
        ['TC-04','Owner enrolls a member with an active plan.','The callable workflow validates the owner/gym relationship and returns enrollment feedback.'],
        ['TC-05','Payment selection is omitted.','The enrollment is recorded with pending payment support.'],
        ['TC-06','Owner searches trainer records.','Matching trainer records are shown without leaving the management screen.'],
        ['TC-07','Member opens today\'s workout and completes a session.','Exercise progress is recorded and the completion screen is reached.'],
        ['TC-08','Owner requests report or finance export.','The export workflow validates selections and handles a returned download result or an error.']
    ],[.65,2.55,3.3])
    heading(doc,'9. FEATURE IMPROVEMENT AND FUTURE ENHANCEMENTS',1)
    add_para(doc,'The items in this section are not claimed as completed functionality. They include the attendance scope moved from the original report, the remaining dashboard work and the trainer-role work requested for future development.')
    heading(doc,'9.1 Attendance for Members and Owners',2); bullets(doc,['Member check-in, attendance status and history.','Owner-side attendance recording, daily/history views and filtering.','Attendance summaries and operational analytics after a real attendance data source is implemented.'])
    heading(doc,'9.2 Remaining Member Dashboard Improvements',2); bullets(doc,['Complete the notification action currently marked for future navigation.','Complete the progress quick action and deepen links between the dashboard, workout and progress areas.','Add any future dashboard indicators only after their supporting data and workflows exist.'])
    heading(doc,'9.3 Remaining Owner Dashboard Improvements',2); bullets(doc,['Complete all dashboard shortcut destinations and "See all" navigation where UI labels exist without a corresponding route action.','Extend dashboard summaries only when the underlying operational data is available.','Continue production validation of report/finance export and storage configuration.'])
    heading(doc,'9.4 Dedicated Trainer Role',2); bullets(doc,['Replace the current TrainerShell placeholder with a role-aware trainer dashboard.','Add trainer-specific navigation and operational screens only after the necessary data model and workflows are implemented.','Keep owner-side trainer enrollment/management separate from the future trainer self-service experience.'])
    heading(doc,'10. CONCLUSION',1)
    add_para(doc,'SweatSync provides a structured foundation for gym operations through connected owner and member workflows. The implemented application includes membership, member, trainer, gym, finance, reporting, profile, workout and progress modules, supported by Firebase services and a feature-first architecture. This design keeps screens, business entities and persistence concerns appropriately separated.')
    add_para(doc,'Attendance, a complete trainer experience and the remaining dashboard actions are correctly treated as future work. Their implementation can build on the existing roles, repositories and Firebase integration without requiring the completed modules to be redesigned.')
    heading(doc,'11. REFERENCES AND BIBLIOGRAPHY',1)
    for ref in ['[1] Flutter Documentation. https://docs.flutter.dev/','[2] Dart Documentation. https://dart.dev/','[3] Firebase Documentation. https://firebase.google.com/docs','[4] Firebase Authentication Documentation. https://firebase.google.com/docs/auth','[5] Cloud Firestore Documentation. https://firebase.google.com/docs/firestore','[6] Firebase Cloud Functions Documentation. https://firebase.google.com/docs/functions','[7] Firebase Cloud Storage Documentation. https://firebase.google.com/docs/storage','[8] Riverpod Documentation. https://riverpod.dev/','[9] GoRouter package documentation. https://pub.dev/packages/go_router']:
        add_para(doc,ref)
    add_page_break(doc); heading(doc,'APPENDIX A. CURRENT IMPLEMENTATION STATUS',1)
    table(doc,['Module','Status','Notes'],[
        ['Authentication and role access','Implemented','Firebase Authentication, role selection and owner/member redirect logic.'],
        ['Owner dashboard','Implemented with remaining actions','Dashboard summary, management links and activity presentation; some shortcuts remain future work.'],
        ['Member dashboard','Implemented with remaining actions','Profile summary, membership, workout, weekly progress, body metrics and quick actions.'],
        ['Gym management','Implemented','Owner can create, view and edit gym information.'],
        ['Member management and enrollment','Implemented','Search/filter/details plus controlled multi-step enrollment.'],
        ['Membership plans','Implemented','Create/edit/activate/deactivate/delete with validation.'],
        ['Trainer management','Implemented owner-side','Trainer records, search/details and enrollment workflow; no completed trainer dashboard.'],
        ['Workout and progress','Implemented','Assigned workouts, sessions, completion records and progress display.'],
        ['Finance and reports','Implemented / deployment dependent','Finance transactions, analytics and export UI; production export depends on Firebase service configuration.'],
        ['Attendance','Future improvement','No attendance module is implemented.'],
    ],[1.7,1.65,3.15])
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__': build()
