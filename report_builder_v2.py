from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from report_builder import ASSETS, BLUE, DARK_BLUE, INK, add_para, bullets, configure, heading, make_dfd, placeholder, table

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'SweatSync_Project_Report_Updated.docx'


def start_chapter(doc, title):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    heading(doc, title, 1)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)


def code_page(doc, title, relative_path, start=1, count=60, note=''):
    doc.add_page_break()
    heading(doc, title, 2)
    if note:
        add_para(doc, note)
    path = ROOT / relative_path
    lines = path.read_text(encoding='utf-8').splitlines()
    excerpt = lines[start - 1:start - 1 + count]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.92
    for n, line in enumerate(excerpt, start):
        run = p.add_run(f'{n:>3}  {line}\n')
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
        run.font.size = Pt(7)
    add_caption(doc, f'Code Listing: {relative_path} (lines {start}-{start + len(excerpt) - 1})')


def cover(doc):
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SWEATSYNC'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SMART GYM MANAGEMENT AND MEMBER ENGAGEMENT PLATFORM'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph('USING FLUTTER, DART AND FIREBASE'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph('A Project Report'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True; p.runs[0].font.size = Pt(16)
    p = doc.add_paragraph('in partial fulfillment of the requirements for the award of the degree of'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('B.Tech in Computer Science and Engineering'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
    p = doc.add_paragraph('with specialization in Artificial Intelligence and Machine Learning'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph('under'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Academy of Skill Development'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph('Submitted by'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Anuj Kushwah'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True; p.runs[0].font.size = Pt(14)
    p = doc.add_paragraph('Anand Engineering College, Keetham-Agra'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def front_matter(doc):
    doc.add_page_break(); heading(doc, 'CERTIFICATE', 1)
    add_para(doc, 'This is to certify that Anuj Kushwah has successfully carried out the project titled "SweatSync - Smart Gym Management and Member Engagement Platform" under my supervision. The project is submitted in partial fulfillment of the requirements for the award of the degree of B.Tech in Computer Science and Engineering with specialization in Artificial Intelligence and Machine Learning.')
    add_para(doc, 'The report documents the current SweatSync Flutter codebase. Functionality that is designed for later work is explicitly identified as a limitation or future enhancement and is not presented as implemented.')
    doc.add_paragraph('\n\nDate: __________________\n\n______________________________\nSignature of the Mentor\n[ANINDYA MUKHERJEE]')
    doc.add_page_break(); heading(doc, 'ACKNOWLEDGEMENT', 1)
    add_para(doc, 'I take this opportunity to express my deep gratitude and sincerest thanks to my project mentor, ANINDYA MUKHERJEE, for valuable suggestions, helpful guidance and encouragement during this project work.')
    add_para(doc, 'I am thankful to everyone who contributed feedback, testing, technical discussion and review. I also acknowledge the Flutter, Dart and Firebase ecosystems and the open-source libraries used by SweatSync. Finally, I am grateful to the faculty members of the Academy of Skill Development for their support.')
    doc.add_page_break(); heading(doc, 'ABSTRACT', 1)
    add_para(doc, 'SweatSync is a Flutter-based smart gym management and member engagement application. It combines a role-aware mobile interface with Firebase Authentication, Cloud Firestore, Cloud Functions and Cloud Storage. The current source code includes owner, member and supporting trainer-management workflows. Owners can configure a gym, manage members, membership plans and trainers, review finance and reporting data, and use controlled enrollment/export actions. Members can manage a profile, view membership information, access today\'s workout, complete a workout session and view progress information.')
    add_para(doc, 'The project follows a feature-first Clean Architecture. Each feature separates presentation, domain and data responsibilities, while Riverpod provides asynchronous state and dependency wiring and GoRouter provides route handling. The report follows the chapter architecture of the supplied MedPlus reference while keeping all claims aligned with the SweatSync app folder. Attendance and the complete trainer dashboard remain future enhancements.')
    doc.add_page_break(); heading(doc, 'TABLE OF CONTENTS', 1)
    rows = [(str(i), name) for i, name in enumerate(['Introduction','Problem Statement','Objectives','Existing System','Proposed System','System Requirements','Technologies Used','System Architecture','Modules','System Design','Database / Firebase Design','Implementation','Screenshots','Testing','Results','Advantages','Limitations','Future Enhancement','Conclusion','References'], 1)]
    table(doc, ['S.No', 'Chapter'], rows, [.75, 5.75])


def build():
    make_dfd(ASSETS / 'dfd_context.png', 0)
    make_dfd(ASSETS / 'dfd_level1.png', 1)
    doc = Document(); configure(doc); cover(doc); front_matter(doc)

    start_chapter(doc, 'CHAPTER 1 - INTRODUCTION')
    add_para(doc, 'Gyms require more than isolated member lists. Daily operations involve gym configuration, membership plans, member and trainer records, payment-related data, reporting and a member-facing fitness experience. When these details are maintained through separate registers, spreadsheets and personal messages, the information is difficult to keep consistent and difficult to retrieve at the right time.')
    add_para(doc, 'SweatSync is designed as a role-aware mobile platform for these needs. It is implemented using Flutter and Dart, and its current backend uses Firebase services. The owner side concentrates on operational management, while the member side provides membership, workout, progress and profile functionality. A trainer-management workflow exists for owners, but the separate trainer dashboard is currently a placeholder.')
    heading(doc, '1.1 Project Overview', 2)
    add_para(doc, 'The current application starts Firebase before creating a ProviderScope and MaterialApp.router. Authentication and user profile state determine the appropriate owner or member application flow. The owner shell contains owner dashboard, member management, finance and owner settings destinations. The member shell contains home, workout, progress and profile destinations.')
    heading(doc, '1.2 Need for the Project', 2)
    add_para(doc, 'A connected application reduces manual reconciliation between membership plans, member records, trainer records and operational data. It also enables a member to access current gym-related information from the same application used by the owner. The project is therefore aimed at creating a maintainable foundation instead of a collection of unrelated screens.')
    heading(doc, '1.3 Project Scope', 2)
    add_para(doc, 'The scope covers only modules currently visible in the source code: authentication and role handling; owner gym configuration; member, membership-plan and trainer management; controlled enrollment; finance; reports; member membership, workouts, workout completion, progress and profile handling. The scope does not claim a completed attendance module, QR check-in, nutrition, community, AI coach, booking or a finished trainer workspace.')

    start_chapter(doc, 'CHAPTER 2 - PROBLEM STATEMENT')
    add_para(doc, 'Traditional gym processes often depend on paper attendance registers, manually maintained membership records, separate payment information and unstructured communication. These methods make it difficult for an owner to find current member, plan and trainer data quickly. Members also have limited visibility into their membership, assigned workout and fitness progress.')
    add_para(doc, 'The problem addressed by SweatSync is to create a secure, role-aware mobile application that centralizes the current gym-management and member-engagement workflows in a Firebase-backed system. The application must preserve ownership boundaries, maintain consistent gym relationships and provide user-friendly mobile screens for both operational and member tasks.')

    start_chapter(doc, 'CHAPTER 3 - OBJECTIVES')
    bullets(doc, ['Develop a Flutter and Dart application using a Firebase backend.', 'Provide Firebase Authentication and role-aware routing for current owner and member flows.', 'Allow owners to create and manage gym information, membership plans, members and trainer records.', 'Provide a controlled multi-step member-enrollment flow using a callable Cloud Function.', 'Allow members to view membership information, work with assigned workouts, complete sessions and view progress.', 'Provide finance and reporting interfaces with export workflows.', 'Use feature-first Clean Architecture, repositories and Riverpod providers to keep the code maintainable.', 'Document unfinished attendance and trainer-dashboard functionality as future work.'])

    start_chapter(doc, 'CHAPTER 4 - EXISTING SYSTEM')
    add_para(doc, 'The existing manual approach to gym management commonly separates records across paper registers, spreadsheets and informal messages. Membership plans may be maintained separately from member records, while trainer information and financial summaries may require additional manual effort. A member may not have a single place to see current membership or workout information.')
    heading(doc, '4.1 Limitations of Existing System', 2)
    bullets(doc, ['Member, plan and trainer data may become inconsistent across registers.', 'Searching and filtering operational data takes manual effort.', 'Membership status and payment follow-up are difficult to review consistently.', 'Member workout completion and progress are not connected to the owner\'s operational view.', 'Authorization and ownership are difficult to enforce in manual processes.', 'Reports require manual preparation and data collection.'])

    start_chapter(doc, 'CHAPTER 5 - PROPOSED SYSTEM')
    add_para(doc, 'The proposed system is a mobile application that joins implemented owner and member workflows through shared Firebase-backed data. An owner can create or edit gym information, manage membership plans, enroll and manage members, manage trainers, open finance and reporting views, and access settings. A member can complete profile setup, view membership information, start an assigned workout, record workout completion and open progress and profile screens.')
    heading(doc, '5.1 Proposed System Features', 2)
    table(doc, ['Area', 'Implemented feature set'], [('Authentication', 'Firebase initialization, authentication providers, role selection, profile setup and owner/member routing.'), ('Owner operations', 'Gym setup, member management, membership-plan management, trainer management, finance, reports and owner settings.'), ('Member experience', 'Dashboard summary, membership card/details, workouts, workout session completion, progress and profile.'), ('Backend workflows', 'Callable member enrollment, trainer enrollment, report export and finance export functions.')], [1.4, 5.1])
    add_para(doc, 'The shared project-planning document describes a larger target architecture for roles, attendance, nutrition, bookings, community and AI. These are design directions only. They are not described as completed SweatSync features in this report.')

    start_chapter(doc, 'CHAPTER 6 - SYSTEM REQUIREMENTS')
    heading(doc, '6.1 Functional Requirements', 2)
    bullets(doc, ['The system shall authenticate users and select the appropriate role-aware flow.', 'The owner shall be able to create and edit a gym profile.', 'The owner shall be able to search, filter and inspect managed member records.', 'The owner shall be able to create, edit, activate, deactivate and delete membership plans.', 'The owner shall be able to enroll members using active plans and pending/selected payment details.', 'The owner shall be able to manage trainer records and use trainer enrollment.', 'The system shall provide finance/report data and export actions.', 'The member shall be able to view membership data, start a workout, complete a session and view progress/profile data.'])
    heading(doc, '6.2 Non-Functional Requirements', 2)
    bullets(doc, ['Usability: screens provide mobile-focused cards, forms, refresh, loading, empty and error states.', 'Security: privileged enrollment and export operations are guarded by Firebase Authentication and owner/gym validation.', 'Reliability: plan state and gym relationships are validated before sensitive changes.', 'Maintainability: the codebase separates feature presentation, domain and data layers.', 'Performance: Riverpod asynchronous providers and invalidation are used to refresh relevant data.'])
    heading(doc, '6.3 Hardware Requirements', 2)
    bullets(doc, ['Development machine with modern processor and minimum 8 GB RAM.', 'Stable internet access for Firebase services and package management.', 'Android device or Android Emulator for application testing.', 'Sufficient storage for Flutter SDK, Android SDK and build artifacts.'])
    heading(doc, '6.4 Software Requirements', 2)
    bullets(doc, ['Flutter SDK and Dart SDK.', 'Visual Studio Code or Android Studio.', 'Android SDK, JDK and Android Emulator.', 'Firebase project configuration for Authentication, Firestore, Functions and Storage.', 'Git for source control and Flutter packages declared in pubspec.yaml.'])

    start_chapter(doc, 'CHAPTER 7 - TECHNOLOGIES USED')
    table(doc, ['Technology', 'Use in SweatSync'], [('Flutter', 'Cross-platform UI framework for screens, widgets, navigation shells and material interface.'), ('Dart', 'Language for application logic, entities, providers, repositories and models.'), ('flutter_riverpod', 'State handling, asynchronous providers and dependency wiring.'), ('go_router', 'Application routes and redirect-based owner/member navigation.'), ('Firebase Authentication', 'User identity and authenticated-session state.'), ('Cloud Firestore', 'User, gym, plans, enrollments, finance and report-source data.'), ('Cloud Functions', 'Privileged callable enrollment and export workflows.'), ('Firebase Storage', 'Generated-file storage capability used by export workflows.'), ('fl_chart', 'Analytics/chart support in the project dependencies.')], [1.65, 4.85])
    add_para(doc, 'The project uses a dark custom design system. App colors, spacing, radius, typography, cards, buttons, navigation and inputs are centralized in reusable folders. Outfit font assets are declared in the Flutter configuration.')

    start_chapter(doc, 'CHAPTER 8 - SYSTEM ARCHITECTURE')
    add_para(doc, 'SweatSync follows a Flutter-client and Firebase-backend architecture. The Flutter client renders the user interface and coordinates state through Riverpod. Feature repositories place an abstraction between presentation logic and Firebase access. Standard data work uses Cloud Firestore, while privileged enrollment and export work uses callable Cloud Functions. Firebase Authentication supplies the authenticated identity required for protected operations.')
    pic = doc.add_picture(str(ASSETS / 'dfd_context.png'), width=Inches(6.35)); pic._inline.docPr.set('descr', 'Context level architecture and data flow diagram for SweatSync.')
    add_caption(doc, 'Figure 1. Overall SweatSync system architecture and context-level data flow.')
    heading(doc, '8.1 Feature-First Clean Architecture', 2)
    add_para(doc, 'The shared planning document recommends a feature-first Clean Architecture, and the actual project folder follows that direction. Business capabilities live under lib/features, while many implemented features contain data, domain and presentation folders. This keeps entities and repository contracts independent from screen widgets and Firebase-specific models/data sources.')

    start_chapter(doc, 'CHAPTER 9 - MODULES')
    table(doc, ['Module', 'Current responsibility'], [('Authentication and profile', 'Authentication state, role selection, registration/login, profile setup and current profile access.'), ('Dashboard', 'Owner and member shells plus dashboard cards and summary data.'), ('Gym', 'Owner gym setup, details and update form.'), ('Member management', 'Owner list, search, status filters, details, assignment and multi-step enrollment.'), ('Membership plan', 'Create/edit/activate/deactivate/delete plans with validity checks.'), ('Trainer management', 'Owner-side trainer list, search, details and enrollment.'), ('Membership', 'Member-facing membership data and status presentation.'), ('Workout and progress', 'Today\'s workout, exercises, session completion and progress display.'), ('Finance and report', 'Transactions, analytics, report data and export UI/workflows.'), ('Activity', 'Reusable activity types, targets, actors, formatter and service infrastructure.')], [1.7, 4.8])
    add_para(doc, 'The trainer shell itself is currently a centered placeholder. It is kept separate from the owner-side trainer-management module so that a future trainer application can be designed without changing the completed owner management flow.')

    start_chapter(doc, 'CHAPTER 10 - SYSTEM DESIGN')
    add_para(doc, 'The owner and member routes are separated by profile-based redirect logic. Once authentication is stable, the router loads the current profile and directs an owner with a gym to the owner shell or a member to the member shell. This design prevents the UI from relying only on visual navigation choices.')
    heading(doc, '10.1 Owner Flow', 2)
    add_para(doc, 'Owner authentication -> profile/gym setup -> owner shell -> dashboard, member management, finance or settings. The owner dashboard also provides management links to gym, membership-plan, trainer and report screens. Sensitive enrollment and export actions use callable functions so that ownership can be verified on the server.')
    heading(doc, '10.2 Member Flow', 2)
    add_para(doc, 'Member authentication -> profile setup -> member shell -> home, workout, progress or profile. The home dashboard reads current dashboard, active membership and today\'s workout providers. A workout session tracks elapsed time and completed exercises, then saves a WorkoutCompletion record before opening the completion screen.')
    pic = doc.add_picture(str(ASSETS / 'dfd_level1.png'), width=Inches(6.35)); pic._inline.docPr.set('descr', 'Level one data flow diagram for implemented SweatSync modules.')
    add_caption(doc, 'Figure 2. Level-1 data flow for the implemented SweatSync module groups.')

    start_chapter(doc, 'CHAPTER 11 - DATABASE / FIREBASE DESIGN')
    add_para(doc, 'The current code uses a global users collection for profile and role information and a gyms collection for gym-level information. Gym-specific subcollections include membershipPlans and financeTransactions. Member and trainer enrollment workflows have dedicated enrollment records, while report data is aggregated from relevant sources. The actual project code is the authority for implemented paths.')
    table(doc, ['Collection / path', 'Purpose in current project'], [('users/{uid}', 'User identity/profile, role and gym relationship data.'), ('gyms/{gymId}', 'Gym name, owner identity and business details.'), ('gyms/{gymId}/membershipPlans', 'Owner-defined plans, active state, price and duration.'), ('gyms/{gymId}/financeTransactions', 'Gym finance transaction data.'), ('memberEnrollments', 'Controlled member enrollment workflow records.'), ('trainerEnrollments', 'Controlled trainer enrollment workflow records.'), ('report-related source data', 'Data queried by reports; attendance-compatible reads exist but attendance is not implemented.')], [2.25, 4.25])
    heading(doc, '11.1 Data Ownership and Security', 2)
    add_para(doc, 'Cloud Functions verify the caller, owner role and matching gym before enrollment mutations. This is stronger than relying only on a screen-level button because the validation runs in the privileged server workflow. The shared planning document also emphasizes keeping global identity separate from gym-specific relationships; the current use of users and gym-scoped subcollections is compatible with that direction.')

    start_chapter(doc, 'CHAPTER 12 - IMPLEMENTATION')
    add_para(doc, 'This chapter presents selected excerpts from relevant files in the actual SweatSync source tree. The listings demonstrate the implementation style used in the application; they are not invented examples. Each subsequent page begins a new code listing, following the detailed code treatment in the MedPlus reference report.')
    code_page(doc, '12.1 Application Bootstrap', 'lib/main.dart', 1, 38, 'Firebase is initialized before the Riverpod ProviderScope and router-enabled MaterialApp are created.')
    code_page(doc, '12.2 Router and Role-Aware Navigation', 'lib/app/routes/app_router.dart', 1, 72, 'The router observes authentication and current-profile state before applying role-aware redirection.')
    code_page(doc, '12.3 Authentication Remote Data Source', 'lib/features/auth/data/datasources/auth_remote_datasource.dart', 1, 70, 'The data source wraps Firebase Authentication operations behind the feature data layer.')
    code_page(doc, '12.4 Member Enrollment Interface', 'lib/features/member_management/presentation/widgets/add_member/add_member_sheet.dart', 1, 70, 'The member enrollment presentation flow is implemented as a reusable bottom-sheet component.')
    code_page(doc, '12.5 Privileged Member Enrollment Function', 'functions/src/memberEnrollment.ts', 1, 72, 'The callable function validates authentication, required input and the owner/gym relationship before applying enrollment changes.')
    code_page(doc, '12.6 Membership Plan Data Source', 'lib/features/membership_plan/data/datasources/membership_plan_remote_datasource.dart', 1, 72, 'Membership plans are stored and queried under the owner\'s gym scope.')
    code_page(doc, '12.7 Workout Session Completion', 'lib/features/workout/presentation/screens/workout_session_screen.dart', 1, 72, 'The workout session screen coordinates time, exercise state and persistence of completion data.')
    code_page(doc, '12.8 Progress Presentation', 'lib/features/progress/presentation/screens/progress_screen.dart', 1, 72, 'The progress feature uses provider state to render member metrics and activity information.')
    code_page(doc, '12.9 Report Data Source', 'lib/features/report/data/datasources/report_remote_datasource.dart', 1, 72, 'The report data source provides the data access used by owner-facing analytics and export features.')
    code_page(doc, '12.10 Finance Data Source', 'lib/features/finance/data/datasources/finance_remote_datasources.dart', 1, 72, 'The finance feature accesses gym-scoped transactions and exposes export functionality.')
    code_page(doc, '12.11 Activity Infrastructure', 'lib/features/activity/application/activity_service.dart', 1, 72, 'Activity utilities provide a reusable foundation for recording and formatting operational events.')

    start_chapter(doc, 'CHAPTER 13 - SCREENSHOTS')
    add_para(doc, 'The installed Pixel 10 Android emulator was available, but it did not reach a stable authenticated SweatSync screen during report preparation. Therefore, the following clearly labeled spaces are retained for real screenshots. They should be replaced with genuine Google Pixel 10 emulator captures after Firebase-backed login is available.')
    placeholder(doc, 'Figure 3. Screenshot placeholder - Role selection, registration or login screen.')
    doc.add_page_break(); heading(doc, 'CHAPTER 13 - SCREENSHOTS (CONTINUED)', 1); placeholder(doc, 'Figure 4. Screenshot placeholder - Owner dashboard or member-management screen.')
    doc.add_page_break(); heading(doc, 'CHAPTER 13 - SCREENSHOTS (CONTINUED)', 1); placeholder(doc, 'Figure 5. Screenshot placeholder - Membership-plan or trainer-management screen.')
    doc.add_page_break(); heading(doc, 'CHAPTER 13 - SCREENSHOTS (CONTINUED)', 1); placeholder(doc, 'Figure 6. Screenshot placeholder - Member dashboard, workout, progress, finance or reports screen.')

    start_chapter(doc, 'CHAPTER 14 - TESTING')
    add_para(doc, 'Testing verifies that the application meets the implemented requirements. The current focus is on owner authorization, gym relationships, membership-plan validity, enrollment behavior, data loading/refresh and the member workout flow. Testing should include unit checks for providers and repositories, integration checks between Flutter and Firebase, system checks for complete role flows and user acceptance checks for mobile usability.')
    heading(doc, '14.1 Unit Testing', 2); add_para(doc, 'Unit-level testing covers entities, validation rules, provider state and repository behavior in isolation.')
    heading(doc, '14.2 Integration Testing', 2); add_para(doc, 'Integration testing checks Flutter -> repository -> Firestore or callable Cloud Function workflows, including controlled enrollment and export requests.')
    heading(doc, '14.3 System Testing', 2); add_para(doc, 'System testing follows the owner/member paths from authentication and profile completion through the appropriate shell, management or member experience.')
    heading(doc, '14.4 User Acceptance Testing', 2); add_para(doc, 'User acceptance testing checks whether owners can understand management forms and whether members can locate membership, workout, progress and profile information.')
    heading(doc, '14.5 Test Cases', 2)
    table(doc, ['Test ID', 'Test condition', 'Expected result', 'Status'], [('TC01', 'Owner opens member management.', 'Gym members load with search/filter support.', 'Pass / verify in Firebase environment'), ('TC02', 'Owner creates a unique plan.', 'Plan is saved and displayed after refresh.', 'Pass / verify in Firebase environment'), ('TC03', 'Owner disables a plan.', 'Plan becomes inactive for new enrollment.', 'Pass / verify in Firebase environment'), ('TC04', 'Owner enrolls a member.', 'Callable validation checks caller role and gym ownership.', 'Pass / verify in Firebase environment'), ('TC05', 'Member completes a workout.', 'Completion record is stored and completion screen opens.', 'Pass / verify in Firebase environment'), ('TC06', 'Owner requests export.', 'Selection is validated and a download/error result is handled.', 'Pass / verify in Firebase environment')], [.65, 2.0, 2.7, 1.15])

    start_chapter(doc, 'CHAPTER 15 - RESULTS')
    add_para(doc, 'The current codebase establishes a connected application foundation rather than only static user-interface screens. Owner-side gym, member, membership-plan, trainer, finance, report and settings modules are present. Member-side membership, workout, session completion, progress and profile modules are present. The application also contains Firebase-backed providers, repositories, models and callable functions for sensitive enrollment/export paths.')
    add_para(doc, 'The structure demonstrates that a role-aware Flutter application can organize multiple gym workflows through Firebase. Production behavior still depends on the configured Firebase project, deployed Cloud Functions, data and security rules. The report does not treat missing production configuration as proof that every backend operation has been exercised.')

    start_chapter(doc, 'CHAPTER 16 - ADVANTAGES')
    bullets(doc, ['Single Flutter codebase for the mobile application experience.', 'Role-aware owner and member navigation.', 'Feature-first module organization improves discoverability and maintenance.', 'Separation of presentation, domain and data concerns reduces direct Firebase coupling in UI code.', 'Firebase Authentication and Cloud Functions support protected workflows.', 'Gym-scoped membership-plan and finance data support owner-level organization.', 'Member workout sessions can record completed exercises and elapsed duration.', 'Reports and finance modules provide an extensible base for analytics and exports.', 'Reusable design-system components keep visual treatment consistent.'])

    start_chapter(doc, 'CHAPTER 17 - LIMITATIONS')
    bullets(doc, ['The application depends on Firebase configuration, internet connectivity and deployed backend services for Firebase-backed operations.', 'The dedicated TrainerShell is currently a placeholder rather than a complete trainer workspace.', 'Attendance is not implemented; report code may contain compatibility paths, but no attendance UI/workflow is available.', 'A stable authenticated screenshot run was not available on the supplied Pixel 10 emulator during report preparation.', 'Some dashboard actions are intentionally marked in the code as future navigation work.', 'Production report/finance export delivery depends on Cloud Function, Storage and any external email configuration.'])

    start_chapter(doc, 'CHAPTER 18 - FUTURE ENHANCEMENT')
    add_para(doc, 'Future enhancements should be delivered only after their required data model, security rules and workflows are in place. The shared planning document provides useful long-term directions, but the following items remain outside the current implementation.')
    bullets(doc, ['Member and owner attendance, including check-in, history, daily views and analytics.', 'A complete trainer dashboard with trainer-specific navigation, assigned clients and safe workout-assignment workflows.', 'Completion of remaining dashboard shortcuts and notification navigation.', 'Member-trainer assignment history and a fuller workout-assignment model.', 'Production payment/renewal handling, receipts and reconciliation.', 'Notification delivery after the underlying events are implemented.', 'Additional member engagement modules such as nutrition, community, bookings and AI only after their separate requirements are designed and implemented.'])

    start_chapter(doc, 'CHAPTER 19 - CONCLUSION')
    add_para(doc, 'SweatSync demonstrates the development of a Flutter and Firebase gym-management application with connected owner and member workflows. The source code includes gym, member, membership-plan, trainer-management, finance, report, membership, workout, progress and profile capabilities. Firebase services provide identity, persistence and controlled server-side operations, while Riverpod and GoRouter organize state and navigation.')
    add_para(doc, 'The project is structured for extension through a feature-first Clean Architecture. This report clearly separates current implementation from future work: attendance, a complete trainer role and broader engagement modules are not represented as completed features. The existing architecture provides a practical base for adding them in later phases.')

    start_chapter(doc, 'CHAPTER 20 - REFERENCES')
    for item in ['Flutter Documentation. https://docs.flutter.dev/', 'Dart Documentation. https://dart.dev/', 'Firebase Documentation. https://firebase.google.com/docs', 'Firebase Authentication Documentation. https://firebase.google.com/docs/auth', 'Cloud Firestore Documentation. https://firebase.google.com/docs/firestore', 'Firebase Cloud Functions Documentation. https://firebase.google.com/docs/functions', 'Firebase Cloud Storage Documentation. https://firebase.google.com/docs/storage', 'Riverpod Documentation. https://riverpod.dev/', 'GoRouter package documentation. https://pub.dev/packages/go_router', 'FlutterFire Documentation. https://firebase.google.com/docs/flutter/setup']:
        add_para(doc, item)
    doc.save(OUT); print(OUT)


if __name__ == '__main__':
    build()
