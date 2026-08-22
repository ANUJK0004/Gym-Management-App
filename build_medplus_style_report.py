from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'SweatSync_Project_Report_Updated.docx'
ASSETS = ROOT / 'report_assets'
ASSETS.mkdir(exist_ok=True)


def font(size=26, bold=False):
    return ImageFont.truetype('C:/Windows/Fonts/arialbd.ttf' if bold else 'C:/Windows/Fonts/arial.ttf', size)


def system_architecture():
    path = ASSETS / 'system_architecture.png'
    im = Image.new('RGB', (1800, 1020), 'white'); d = ImageDraw.Draw(im)
    def box(x, y, w, h, text, fill='E8EEF5'):
        d.rounded_rectangle((x,y,x+w,y+h), radius=24, fill='#'+fill, outline='#1F4D78', width=4)
        lines = text.split('\n'); yy = y + (h - len(lines)*36)//2
        for line in lines:
            b=d.textbbox((0,0),line,font=font(28)); d.text((x+(w-(b[2]-b[0]))//2,yy),line,font=font(28),fill='#0B2545'); yy += 38
    def arrow(a,b,label):
        d.line((a[0],a[1],b[0],b[1]),fill='#455A64',width=4); d.polygon([(b[0],b[1]),(b[0]-20,b[1]-10),(b[0]-20,b[1]+10)],fill='#455A64')
        bb=d.textbbox((0,0),label,font=font(20)); x=(a[0]+b[0]-(bb[2]-bb[0]))//2; y=(a[1]+b[1]-(bb[3]-bb[1]))//2-24; d.rectangle((x-4,y-2,x+(bb[2]-bb[0])+4,y+(bb[3]-bb[1])+2),fill='white'); d.text((x,y),label,font=font(20),fill='#37474F')
    d.text((530,24),'Figure 1. SweatSync System Architecture',font=font(38,True),fill='#0B2545')
    box(80,370,250,150,'OWNER\nMEMBER\nTRAINER')
    box(570,330,380,220,'FLUTTER APPLICATION\nPresentation Layer\nRiverpod + GoRouter','DCEAF7')
    box(1210,120,300,130,'FIREBASE AUTH','F4F6F9')
    box(1210,340,300,130,'CLOUD FIRESTORE','F4F6F9')
    box(1210,560,300,130,'CLOUD FUNCTIONS\n& STORAGE','F4F6F9')
    arrow((330,445),(570,445),'role-aware UI actions')
    arrow((950,390),(1210,185),'sign-in / session')
    arrow((950,440),(1210,405),'read / write data')
    arrow((950,500),(1210,625),'privileged operations / exports')
    im.save(path)
    return path


def set_cell(cell, fill=None):
    tcPr = cell._tc.get_or_add_tcPr()
    mar=OxmlElement('w:tcMar')
    for side, value in [('top','80'),('start','120'),('bottom','80'),('end','120')]:
        e=OxmlElement('w:'+side); e.set(qn('w:w'),value); e.set(qn('w:type'),'dxa'); mar.append(e)
    tcPr.append(mar)
    if fill:
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); h=OxmlElement('w:tblHeader'); h.set(qn('w:val'),'true'); trPr.append(h)


def set_widths(table, widths):
    pr=table._tbl.tblPr; layout=OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); pr.append(layout)
    for row in table.rows:
        for cell,width in zip(row.cells,widths): cell.width=Inches(width)


def add_field(p, code):
    r=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text=code; s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); t=OxmlElement('w:t'); t.text='1'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r._r.extend([b,i,s,t,e])


def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(.85); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    n=doc.styles['Normal']; n.font.name='Calibri'; n._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); n._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(5); n.paragraph_format.line_spacing=1.08
    for name,size,col,bef,aft in [('Title',24,'0B2545',0,8),('Heading 1',16,'2E74B5',0,10),('Heading 2',13,'2E74B5',8,6),('Heading 3',11.5,'1F4D78',6,4)]:
        s=doc.styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft)
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('SweatSync Project Report | '); add_field(f,'PAGE')


def new_page(doc): doc.add_page_break()


def chapter(doc, title):
    new_page(doc); p=doc.add_heading(title,1); p.paragraph_format.keep_with_next=True


def para(doc, text, italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.08; r=p.add_run(text); r.italic=italic; return p


def bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(item)


def table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; set_repeat_header(t.rows[0])
    for c,h in zip(t.rows[0].cells,headers):
        c.text=h; set_cell(c,'F2F4F7')
        for r in c.paragraphs[0].runs: r.bold=True
    for vals in rows:
        cells=t.add_row().cells
        for c,v in zip(cells,vals): c.text=v; set_cell(c)
    set_widths(t,widths); doc.add_paragraph()


def code_block(doc, title, path, start=1, count=48):
    para(doc, title)
    lines=(ROOT/path).read_text(encoding='utf-8').splitlines()[start-1:start-1+count]
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; set_repeat_header(t.rows[0]); c=t.cell(0,0); set_cell(c,'FAFAFA')
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    for idx,line in enumerate(lines,start):
        r=p.add_run(f'{idx:>3}  {line}\n'); r.font.name='Consolas'; r._element.rPr.rFonts.set(qn('w:ascii'),'Consolas'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Consolas'); r.font.size=Pt(7.1)
    cap=doc.add_paragraph(f'Code Listing: {Path(path).as_posix()} (lines {start}-{start+len(lines)-1})'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs: r.italic=True; r.font.size=Pt(8.5)


def placeholder(doc, label):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; set_repeat_header(t.rows[0]); c=t.cell(0,0); set_cell(c,'FAFAFA')
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    r=p.add_run('\n\n\n\n\nINSERT GOOGLE PIXEL 10 EMULATOR SCREENSHOT\n\n\n\n\n'); r.bold=True; r.font.color.rgb=RGBColor(105,105,105); r.font.size=Pt(10)
    p=doc.add_paragraph(label); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.italic=True; r.font.size=Pt(8.5)


def build():
    arch=system_architecture()
    from report_builder import make_dfd
    make_dfd(ASSETS/'dfd_context.png',0); make_dfd(ASSETS/'dfd_level1.png',1)
    d=Document(); setup(d)
    # Pages 1-5
    for _ in range(4): d.add_paragraph()
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SWEATSYNC'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string('0B2545')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SMART GYM MANAGEMENT AND MEMBER ENGAGEMENT PLATFORM'); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string('2E74B5')
    p=d.add_paragraph('USING FLUTTER, DART AND FIREBASE'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4): d.add_paragraph()
    p=d.add_paragraph('A Project Report'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(16); p.runs[0].bold=True
    for text in ['in partial fulfillment of the requirements for the award of the degree of','B.Tech in Computer Science and Engineering','with specialization in Artificial Intelligence and Machine Learning','under','Academy of Skill Development','Submitted by','Anuj Kushwah','Anand Engineering College, Keetham-Agra']:
        p=d.add_paragraph(text); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if text in ['B.Tech in Computer Science and Engineering','Academy of Skill Development','Anuj Kushwah']: p.runs[0].bold=True
    new_page(d); d.add_heading('CERTIFICATE',1); para(d,'This is to certify that Anuj Kushwah has successfully carried out the project titled "SweatSync - Smart Gym Management and Member Engagement Platform" under my supervision. The work is submitted in partial fulfillment of the requirements for the award of the degree of B.Tech in Computer Science and Engineering with specialization in Artificial Intelligence and Machine Learning.')
    para(d,'The report documents the present Flutter application, its Firebase-backed architecture and the implemented owner and member workflows. Planned features are explicitly separated in the Future Enhancement chapter.')
    d.add_paragraph('\n\nDate: __________________\n\n______________________________\nSignature of the Mentor\n[ANINDYA MUKHERJEE]')
    new_page(d); d.add_heading('ACKNOWLEDGMENT',1); para(d,'I take this opportunity to express my deep gratitude and sincerest thanks to my project mentor, ANINDYA MUKHERJEE, for valuable suggestions, helpful guidance and encouragement during the execution of this project work. I am grateful to all contributors who provided review, feedback and technical discussion. I also acknowledge the Flutter, Dart and Firebase ecosystems and the faculty members of the Academy of Skill Development for their support.')
    new_page(d); d.add_heading('ABSTRACT',1); para(d,'SweatSync is a Flutter mobile application for structured gym management and member engagement. The system uses Firebase Authentication for identity, Cloud Firestore for data, Cloud Functions for controlled server workflows and Cloud Storage for generated export files. It supports role-aware owner and member experiences, with owner-side trainer enrollment and management.')
    para(d,'Implemented owner capabilities include gym management, dashboard summaries, members, membership plans, member enrollment, trainers, finance, reports and settings. Implemented member capabilities include profile setup and editing, membership details, workout access, timed workout sessions, workout completion and progress views. The codebase follows a feature-first clean architecture with presentation, domain and data responsibilities separated inside feature modules. Attendance and the complete trainer dashboard are future work, not completed functionality.')
    new_page(d); d.add_heading('TABLE OF CONTENTS',1)
    toc=['1  Introduction','2  Problem Statement','3  Objectives','4  Existing System','5  Proposed System','6  System Requirements','7  Technologies Used','8  System Architecture','9  Modules','10  System Design','11  Database / Firebase Design','12  Implementation','13  Screenshots','14  Testing','15  Results','16  Advantages','17  Limitations','18  Future Enhancement','19  Conclusion','20  References']
    for i,item in enumerate(toc,1): para(d,f'{item}............................................................ {5+i}')
    # Pages 6-11 chapters 1-6
    chapter(d,'CHAPTER 1 - INTRODUCTION'); para(d,'SweatSync is a mobile-first gym management and member engagement platform developed with Flutter and Dart. The project addresses the need to coordinate gym operations such as member records, membership plans, trainers, finance and reports while also providing members with personal membership, workout and progress experiences.')
    d.add_heading('1.1 Project Overview',2); para(d,'The application begins with role selection and authentication, then uses profile data to route users into owner or member experiences. Owners work with gym-scoped information; members work with their profile, membership and assigned workout information. Firebase-backed repositories provide the data layer while Riverpod providers coordinate asynchronous UI state.')
    d.add_heading('1.2 Need for the Project',2); para(d,'Paper registers, separate spreadsheets and manual follow-up make gym records difficult to maintain. SweatSync provides a structured interface for current operational workflows and protects sensitive enrollment actions through server-side validation.')
    chapter(d,'CHAPTER 2 - PROBLEM STATEMENT'); para(d,'Gym owners need a reliable way to manage members, membership plans, trainer records, enrollment payments and operational information. Members need a clear mobile interface for their profile, membership status, assigned workout and progress information. Traditional fragmented processes make it difficult to maintain consistency across these activities.')
    para(d,'The problem addressed by SweatSync is the creation of a role-aware Flutter application that centralizes the implemented gym workflows and connects them to persistent Firebase services without placing sensitive enrollment logic directly in the client interface.')
    chapter(d,'CHAPTER 3 - OBJECTIVES'); bullets(d,['Develop a Flutter and Dart application with role-aware owner and member experiences.','Use Firebase Authentication for user identity and Cloud Firestore for persistent gym data.','Support gym management, member management, membership-plan management and member enrollment.','Provide owner-side trainer management and a controlled trainer-enrollment workflow.','Provide member membership, workout, workout-completion, progress and profile experiences.','Provide finance and reporting screens with export workflows.','Maintain a feature-first clean architecture using Riverpod and GoRouter.'])
    chapter(d,'CHAPTER 4 - EXISTING SYSTEM'); para(d,'A conventional gym commonly maintains membership, trainer and payment information through paper records, spreadsheets or disconnected messaging. Member access to current membership status, assigned workouts and recorded progress can be limited or unavailable. Operational reporting frequently requires manual consolidation.')
    d.add_heading('4.1 Limitations of Existing System',2); bullets(d,['Member and plan records can become inconsistent across sources.','Searching and filtering operational data requires manual effort.','Enrollment and payment information may not be connected to membership records.','Members may not have a single mobile place to see membership and workout information.','Reports and finance records are harder to aggregate without a structured data source.'])
    chapter(d,'CHAPTER 5 - PROPOSED SYSTEM'); para(d,'SweatSync provides a Firebase-backed mobile application with distinct owner and member application experiences. Owner actions are scoped to a gym and include management, enrollment, finance and reports. Member actions focus on membership, workouts, recorded completion, progress and profile information.')
    d.add_heading('5.1 Implemented Proposed-System Features',2); bullets(d,['Firebase email/password authentication with role-aware profile checks.','Gym, membership plan, member and trainer management.','Multi-step member enrollment with membership-plan and payment-pending support.','Owner finance, report dashboards and export workflows.','Member dashboard, membership details, workout session, completion and progress views.','Reusable dark design system with cards, inputs, buttons and navigation.'])
    chapter(d,'CHAPTER 6 - SYSTEM REQUIREMENTS'); d.add_heading('6.1 Functional Requirements',2); bullets(d,['The system shall authenticate users and show a role-aware experience.','Owners shall manage gym data, members, plans, trainers, finance and reports.','Members shall view membership, workout, progress and profile data.','The system shall validate owner/gym context for privileged enrollment workflows.','The system shall record a completed workout and update progress-related records.'])
    d.add_heading('6.2 Non-Functional Requirements',2); bullets(d,['Usability: mobile layouts with loading, empty, error and refresh states.','Security: authenticated and authorized callable functions for sensitive owner actions.','Maintainability: feature-first separation of presentation, domain and data layers.','Scalability: Firebase services and gym-scoped collections support growing operational data.'])
    d.add_heading('6.3 Hardware and Software Requirements',2); table(d,['Category','Requirement'],[['Development computer','Windows 10 or later, 8 GB RAM minimum and stable internet access.'],['Testing device','Android device or Android Emulator; Pixel 10 AVD profiles are available.'],['Development tools','Flutter SDK, Dart SDK, Android Studio or VS Code, Android SDK and Git.'],['Backend','Firebase project setup for Authentication, Firestore, Cloud Functions and Storage.']],[1.55,4.85])
    # 7 two pages
    chapter(d,'CHAPTER 7 - TECHNOLOGIES USED'); d.add_heading('7.1 Flutter and Dart',2); para(d,'Flutter is the cross-platform UI framework used to construct the mobile interface, and Dart is the language used for models, widgets, providers, repositories and Firebase operations. The project uses Material widgets together with a custom dark design system and Outfit fonts.')
    d.add_heading('7.2 Riverpod and GoRouter',2); para(d,'flutter_riverpod is used for provider-based state management, dependency wiring and asynchronous UI data. GoRouter defines routes and redirects users according to authentication and profile state.')
    d.add_heading('7.3 Firebase Authentication and Cloud Firestore',2); para(d,'Firebase Authentication maintains user identity. Cloud Firestore holds users, gyms, plans, finance transactions, enrollment records, workout completion and progress-related documents.')
    new_page(d); d.add_heading('CHAPTER 7 - TECHNOLOGIES USED (CONTINUED)',1); d.add_heading('7.4 Cloud Functions and Cloud Storage',2); para(d,'Callable Cloud Functions are used for privileged member enrollment, trainer enrollment and export operations. Cloud Storage supports generated report files and future file workflows. The client receives validated results from callable functions rather than directly performing sensitive backend mutations.')
    d.add_heading('7.5 Other Declared Packages',2); table(d,['Area','Packages / use'],[['Charts and visuals','fl_chart for chart presentation.'],['Connectivity and secure storage','connectivity_plus and flutter_secure_storage.'],['Scanning and links','mobile_scanner and url_launcher.'],['Notifications packages','firebase_messaging and flutter_local_notifications are declared; notification workflows are not documented as implemented screens.']],[2.0,4.4])
    # 8 two pages
    chapter(d,'CHAPTER 8 - SYSTEM ARCHITECTURE'); para(d,'SweatSync follows a Flutter client and Firebase backend architecture. The Flutter application uses feature modules. Within a mature feature, presentation handles screens and providers, domain holds entities and repository interfaces, and data contains Firebase data sources, models and repository implementations.')
    pic=d.add_picture(str(arch),width=Inches(6.6)); pic._inline.docPr.set('descr','Overall system architecture diagram for SweatSync.'); p=d.add_paragraph('Figure 1: Overall system architecture.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    new_page(d); d.add_heading('CHAPTER 8 - SYSTEM ARCHITECTURE (CONTINUED)',1); d.add_heading('8.1 General Data Flow',2); para(d,'A user interacts with a Flutter screen. Riverpod providers coordinate the action and call a repository. Repository implementations use Firestore or callable Cloud Functions. Standard dashboard and list data is read from Firestore; sensitive enrollment and export actions use callable functions. The returned result updates provider state and the screen refreshes.')
    pic=d.add_picture(str(ASSETS/'dfd_context.png'),width=Inches(6.45)); pic._inline.docPr.set('descr','Context-level data flow diagram for SweatSync.'); p=d.add_paragraph('Figure 2: Context-level data flow diagram.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # 9 three pages
    chapter(d,'CHAPTER 9 - MODULES'); d.add_heading('9.1 Authentication and Role Access',2); para(d,'The auth feature offers sign-up, sign-in, sign-out, password-reset and email-verification actions. A user profile contains role and completion information. Router redirects use the authentication stream and the current user profile to route owners or members appropriately.')
    d.add_heading('9.2 Owner and Gym Module',2); para(d,'The owner setup and gym modules allow an owner to create, display and update gym information. The owner shell provides bottom navigation for dashboard, members, finance and settings.')
    d.add_heading('9.3 Member and Membership-Plan Modules',2); para(d,'Member management provides lists, search, status filters and details. Membership-plan management supports creation, editing, activation/deactivation and deletion with membership-usage validation.')
    new_page(d); d.add_heading('CHAPTER 9 - MODULES (CONTINUED)',1); d.add_heading('9.4 Member Enrollment Module',2); para(d,'The member enrollment experience is a multi-step bottom sheet for personal information, membership-plan selection, fitness goal, start date and payment selection. The callable enrollment function validates the authenticated owner and gym relationship before creating the enrollment record. Existing member accounts can be linked in an atomic update with a finance transaction.')
    d.add_heading('9.5 Trainer Management Module',2); para(d,'Owners can list, search, filter and inspect trainer records. Trainer enrollment is handled through a controlled backend operation. The dedicated trainer shell currently displays only a placeholder, so a standalone trainer dashboard is not considered implemented.')
    d.add_heading('9.6 Finance and Report Modules',2); para(d,'Finance displays revenue-oriented information, breakdowns and transactions. Reports display summary metrics, trends, peak-hours information, membership breakdown and downloadable report entries. Both areas include export flows that call Cloud Functions.')
    new_page(d); d.add_heading('CHAPTER 9 - MODULES (CONTINUED)',1); d.add_heading('9.7 Member Dashboard, Workout and Progress',2); para(d,'The member home screen combines profile summary, membership, today\'s workout, weekly progress, body metrics and quick actions. The workout feature provides workout details, session progress, exercise completion and a completion screen. Progress provides metrics, weekly activity and personal records.')
    d.add_heading('9.8 Profile and Settings',2); para(d,'Profile screens support member data setup, display and editing. Owner settings provide business/header information, administrative account entries, membership-plan settings and logout controls.')
    # 10 two pages
    chapter(d,'CHAPTER 10 - SYSTEM DESIGN'); d.add_heading('10.1 Role and Navigation Flow',2); para(d,'Application start proceeds to splash, role selection and authentication. When a user signs in, the router checks Firebase Authentication state, then reads the user profile. A completed owner profile with a gym is directed to OwnerShell. A completed member profile is directed to MemberShell. The trainer route exists but its UI remains a placeholder.')
    d.add_heading('10.2 Member Workout Flow',2); para(d,'Member dashboard -> Today\'s Workout -> Workout Details or Start Workout -> Workout Session -> Exercise completion -> Workout Completed -> progress-related refresh. The workout session tracks elapsed time and completed exercises. The completion flow writes a completion record and refreshes workout providers.')
    new_page(d); d.add_heading('CHAPTER 10 - SYSTEM DESIGN (CONTINUED)',1); d.add_heading('10.3 Owner Enrollment Flow',2); para(d,'Owner -> Member Management -> Add Member -> Personal Information -> Active Membership Plan -> Payment Selection -> Callable createMemberEnrollment -> owner/gym validation -> enrollment result. When the account already exists and passes validation, the function updates membership details, completes the enrollment and creates a finance transaction in one batch.')
    d.add_heading('10.4 Report and Finance Export Flow',2); para(d,'Owner -> Finance or Reports -> select export options -> callable export function -> generated result / download URL -> application opens the URL. Production operation depends on Firebase project, Cloud Storage and any configured delivery services.')
    # 11 three pages
    chapter(d,'CHAPTER 11 - DATABASE / FIREBASE DESIGN'); d.add_heading('11.1 Authentication and User Profile',2); para(d,'Firebase Authentication provides the authenticated UID. The users collection stores profile data such as role, email, gym relationship, profile completion and timestamps. Role and completion information are used by the application router to make navigation decisions.')
    table(d,['Collection / path','Purpose'],[['users/{uid}','Owner, member and trainer profiles; role and gym relationship.'],['gyms/{gymId}','Gym information owned and edited by the owner.'],['gyms/{gymId}/membershipPlans','Owner-defined membership plans.'],['memberEnrollments','Controlled enrollment records.'],['trainerEnrollments','Trainer enrollment records.']],[2.35,4.05])
    new_page(d); d.add_heading('CHAPTER 11 - DATABASE / FIREBASE DESIGN (CONTINUED)',1); d.add_heading('11.2 Finance and Reporting Data',2); table(d,['Collection / path','Purpose'],[['gyms/{gymId}/financeTransactions','Income/expense transaction data for the gym.'],['npsResponses','Potential NPS source consumed defensively by reports.'],['attendance or gyms/{gymId}/attendance','Read paths currently considered by report data access; no attendance feature is implemented.'],['Cloud Storage','Generated export files and future media/file workflows.']],[2.35,4.05])
    para(d,'The report data source reads members, trainers, plans and finance transactions by gym context. Attendance access is defensive and data-driven only; it does not create a completed attendance feature.')
    new_page(d); d.add_heading('CHAPTER 11 - DATABASE / FIREBASE DESIGN (CONTINUED)',1); d.add_heading('11.3 Workout Completion and Progress Data',2); para(d,'Workout completion is stored in a user-scoped completed collection. A transaction updates the completion document, the user progress summary and weekly activity. Personal-record documents are updated when an exercise weight is greater than the previously stored record. This design prevents the same workout from being recorded twice on the same day for the same workout ID.')
    table(d,['Path','Role in the feature'],[['workout_completions/{userId}/completed/{workoutId_date}','Completed-workout record.'],['progress/{userId}','Aggregate workout/progress values.'],['progress/{userId}/weekly_activity/{day}','Daily activity values.'],['progress/{userId}/personal_records/{exerciseId}','Personal record for a weighted exercise.']],[2.9,3.5])
    # 12 nine pages
    chapter(d,'CHAPTER 12 - IMPLEMENTATION'); para(d,'The source tree is organized by feature and layer. Each current feature keeps related screens, providers, entities, repositories and Firebase implementations close together. The following excerpts are from the current SweatSync source files and demonstrate actual implementation patterns; they are included as concise representative code listings, following the reference report style.')
    table(d,['Layer','Examples'],[['Presentation','screens, widgets and Riverpod providers'],['Domain','entities and repository contracts'],['Data','data sources, models and repository implementations'],['Backend','Firebase callable functions for guarded workflows']],[1.45,4.95])
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.1 Application bootstrap and ProviderScope','lib/main.dart',1,45)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.2 Authentication provider and role validation','lib/features/auth/presentation/providers/auth_provider.dart',1,50)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.3 Role-aware router redirect logic','lib/app/routes/app_router.dart',34,55)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.4 Membership-plan repository implementation','lib/features/membership_plan/data/repositories/membership_plan_repository_impl.dart',1,50)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.5 Callable member-enrollment validation','functions/src/memberEnrollment.ts',1,52)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.6 Member enrollment atomic update and finance record','functions/src/memberEnrollment.ts',125,52)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.7 Workout completion transaction','lib/features/progress/data/services/workout_completion_service.dart',1,55)
    new_page(d); d.add_heading('CHAPTER 12 - IMPLEMENTATION (CONTINUED)',1); code_block(d,'12.8 Finance export callable invocation','lib/features/finance/data/datasources/finance_remote_datasources.dart',1,60)
    # 13 four
    chapter(d,'CHAPTER 13 - SCREENSHOTS'); para(d,'The installed Pixel 10 emulator could not reach a stable authenticated app screen while this report was prepared. The labeled areas below preserve the report layout for real application screenshots. Insert captured screenshots after signing in to a configured Firebase environment.')
    placeholder(d,'Figure 3: Role selection, login or registration screen.')
    placeholder(d,'Figure 4: Member dashboard screen.')
    new_page(d); d.add_heading('CHAPTER 13 - SCREENSHOTS (CONTINUED)',1); placeholder(d,'Figure 5: Member membership, workout or workout-session screen.'); placeholder(d,'Figure 6: Progress or profile screen.')
    new_page(d); d.add_heading('CHAPTER 13 - SCREENSHOTS (CONTINUED)',1); placeholder(d,'Figure 7: Owner dashboard screen.'); placeholder(d,'Figure 8: Member management or member-enrollment screen.')
    new_page(d); d.add_heading('CHAPTER 13 - SCREENSHOTS (CONTINUED)',1); placeholder(d,'Figure 9: Membership-plan or trainer-management screen.'); placeholder(d,'Figure 10: Finance or reports screen.')
    # 14 2
    chapter(d,'CHAPTER 14 - TESTING'); para(d,'Testing verifies that the application behaves according to the implemented requirements. The main focus is authentication and role access, member and membership-plan operations, controlled enrollment, trainer management, workouts, progress, finance and report/export workflows.')
    d.add_heading('14.1 Testing Types',2); bullets(d,['Unit testing for models, validation paths and provider behavior.','Integration testing for Flutter -> repository -> Firebase flows.','System testing for connected owner and member workflows.','User acceptance checks for understandable forms, feedback states and navigation.'])
    new_page(d); d.add_heading('CHAPTER 14 - TESTING (CONTINUED)',1); d.add_heading('14.2 Representative Test Cases',2); table(d,['ID','Test condition','Expected result','Status'],[['TC01','Sign in with a role that matches the profile.','Role-aware entry succeeds.','Pass / configuration dependent'],['TC02','Owner creates a unique membership plan.','Plan is stored and list refreshes.','Pass / configuration dependent'],['TC03','Owner disables a plan.','Plan is inactive for new enrollment.','Pass / configuration dependent'],['TC04','Enroll an existing member under the owner\'s gym.','Callable function validates owner and creates linked records.','Pass / configuration dependent'],['TC05','Member completes a workout.','Completion/progress writes run and UI can refresh.','Pass / configuration dependent'],['TC06','Open finance/report export.','Selection is validated and callable result is handled.','Pass / configuration dependent']],[.5,2.15,2.65,1.15])
    # 15-20 six pages
    chapter(d,'CHAPTER 15 - RESULTS'); para(d,'The current implementation provides connected owner and member workflows rather than isolated screens. Owners can manage gym-scoped members, plans, trainers, finance and reports. Members can access profile, membership, workouts, workout-session completion and progress. Firebase Authentication, Firestore and callable functions provide the core persistence and controlled backend behavior. Report and finance exports depend on deployment and Firebase project configuration for full production operation.')
    chapter(d,'CHAPTER 16 - ADVANTAGES'); bullets(d,['Single Flutter codebase for the application interface.','Feature-first architecture makes related files easier to locate and maintain.','Role-aware routing helps separate owner and member experiences.','Firestore-backed data supports persistent gym, member, plan and progress data.','Callable functions protect sensitive enrollment and export workflows.','Workout completion updates related progress data transactionally.','Reusable design-system components promote consistent dark UI styling.'])
    chapter(d,'CHAPTER 17 - LIMITATIONS'); bullets(d,['Attendance is not implemented and is not claimed as a current feature.','The trainer shell is a placeholder; a dedicated trainer dashboard is not complete.','Some dashboard actions are marked for future navigation.','Report and finance export production delivery depends on deployed Cloud Functions, Storage and external configuration.','The application depends on Firebase connectivity for cloud-backed operations.','Notification packages are declared, but a complete notification workflow is not documented as an implemented feature.'])
    chapter(d,'CHAPTER 18 - FUTURE ENHANCEMENT'); para(d,'The product-planning document describes a larger roadmap. Consistent with the actual codebase, the following items are future work rather than completed modules.')
    bullets(d,['Member and owner attendance: QR/check-in, history, attendance rates and operational analytics.','Dedicated trainer role: trainer dashboard, assigned-member access, workout creation/assignment and progress monitoring.','Remaining member and owner dashboard actions, including notifications and deeper navigation.','Invitation acceptance and secure onboarding for newly enrolled members and trainers.','Optional nutrition, booking, communication, inventory, marketing and community features only after their supporting data and role workflows are implemented.','Production report/finance export delivery, audit history and deeper analytics.'])
    chapter(d,'CHAPTER 19 - CONCLUSION'); para(d,'SweatSync demonstrates a practical Flutter and Firebase implementation for connected gym management and member engagement. The project combines role-aware navigation, a feature-first clean architecture, Firestore persistence and callable backend operations. Current functionality covers gym, member, membership-plan, trainer-management, finance, reports, profile, workout and progress modules. Future attendance, a completed trainer experience and remaining dashboard improvements can extend this foundation without misrepresenting unfinished features as complete.')
    chapter(d,'CHAPTER 20 - REFERENCES');
    refs=['Flutter Documentation. https://docs.flutter.dev/','Dart Documentation. https://dart.dev/','Firebase Documentation. https://firebase.google.com/docs','Firebase Authentication Documentation. https://firebase.google.com/docs/auth','Cloud Firestore Documentation. https://firebase.google.com/docs/firestore','Firebase Cloud Functions Documentation. https://firebase.google.com/docs/functions','Firebase Cloud Storage Documentation. https://firebase.google.com/docs/storage','Riverpod Documentation. https://riverpod.dev/','GoRouter package documentation. https://pub.dev/packages/go_router','Material Design 3. https://m3.material.io/']
    for i,r in enumerate(refs,1): para(d,f'{i}. {r}')
    d.save(OUT); print(OUT)

if __name__=='__main__': build()
